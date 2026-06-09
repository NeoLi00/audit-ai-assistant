from pathlib import Path

from docx import Document as DocxDocument

from app.services.parser.base import ParsedBlock, ParseResult
from app.services.parser.libreoffice_converter import convert_doc_to_docx


class WordParser:
    def parse(self, path: Path) -> ParseResult:
        try:
            parse_path = convert_doc_to_docx(path) if path.suffix.lower() == ".doc" else path
            doc = DocxDocument(parse_path)
            blocks: list[ParsedBlock] = []
            heading_stack: list[str] = []
            para_index = 0
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                style_name = paragraph.style.name if paragraph.style else ""
                if style_name.startswith("Heading") or style_name.startswith("标题"):
                    level = _heading_level(style_name)
                    heading_stack = heading_stack[: level - 1] + [text]
                    block_type = "heading"
                else:
                    block_type = "paragraph"
                blocks.append(
                    ParsedBlock(
                        text=text,
                        block_type=block_type,
                        heading_path="/".join(heading_stack) if heading_stack else None,
                        paragraph_index=para_index,
                    )
                )
                para_index += 1

            for table_index, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    rows.append(" | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells))
                if rows:
                    blocks.append(
                        ParsedBlock(
                            text="\n".join(rows),
                            block_type="table",
                            heading_path="/".join(heading_stack) if heading_stack else None,
                            paragraph_index=table_index,
                        )
                    )
            return ParseResult(status="ready", blocks=blocks, text="\n".join(b.text for b in blocks))
        except Exception as exc:
            return ParseResult(status="failed", error_message=str(exc))


def _heading_level(style_name: str) -> int:
    digits = "".join(ch for ch in style_name if ch.isdigit())
    return max(1, min(int(digits or "1"), 6))

